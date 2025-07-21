#!/usr/bin/env python3
import sys
import os
from docx import Document

def extract_text_from_docx(file_path):
    """Extract text from DOCX file with better error handling"""
    try:
        doc = Document(file_path)
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # Join all text parts
        full_text = "\n".join(text_parts)
        
        # Log extraction success
        print(f"Successfully extracted {len(full_text)} characters from DOCX", file=sys.stderr)
        
        return full_text
        
    except Exception as e:
        print(f"Error extracting DOCX: {str(e)}", file=sys.stderr)
        return ""

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <file_path>", file=sys.stderr)
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    # Extract text
    text = extract_text_from_docx(file_path)
    
    # Output to stdout
    print(text)